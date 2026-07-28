"""
Resume tailoring using Groq API (free tier — works in Ireland/EU).

Get a free API key at: console.groq.com -> API Keys -> Create API Key
Sign up with Google — no billing ever required.
Free limits: 14,400 requests/day.

Only modifies existing content — never adds fake experience or skills.
"""

import json
import logging
import os
import re
import shutil

from docx import Document

logger = logging.getLogger(__name__)


def _extract_resume_text(docx_path: str) -> str:
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text and text not in paragraphs:
                        paragraphs.append(text)
    return "\n".join(paragraphs)


def _apply_replacements(docx_path: str, replacements: list[dict], output_path: str):
    shutil.copy2(docx_path, output_path)
    doc = Document(output_path)

    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            for r in replacements:
                old, new = r.get("old", ""), r.get("new", "")
                if old and new and old in para.text:
                    for run in para.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    doc.save(output_path)


TAILOR_PROMPT = """You are helping tailor a job applicant's resume for a specific role.

JOB TITLE: {job_title}
COMPANY: {company}

JOB DESCRIPTION:
{job_description}

CURRENT RESUME:
{resume_text}

Your task: Suggest specific text replacements to better align the resume with this job.

STRICT RULES:
1. NEVER add skills, experience, qualifications, or achievements that don't already exist in the resume
2. Only rephrase existing content using keywords from the job description where they naturally fit
3. Only modify the professional summary/objective and bullet point descriptions
4. Do not change job titles, company names, dates, or education details
5. Replacements must sound natural and authentic — not keyword-stuffed
6. Make no more than 5-8 targeted changes

Respond ONLY with a JSON array of replacement objects, like this:
[
  {{"old": "exact text from resume to replace", "new": "improved version using job-relevant language"}},
  ...
]

If the resume already aligns well with the job, return an empty array: []
Do not include any explanation or other text — only valid JSON."""


def tailor_resume(
    resume_path: str,
    job_title: str,
    company: str,
    job_description: str,
    output_dir: str,
    api_key: str,
    model: str = "llama-3.1-8b-instant",
) -> str:
    """
    Tailor the resume for a specific job using Groq. Returns path to tailored resume.
    Falls back to original resume if API key is missing or call fails.
    """
    if not api_key or api_key.startswith("YOUR_"):
        logger.warning("No Groq API key configured — using original resume")
        return resume_path

    try:
        from groq import Groq
    except ImportError:
        logger.error("groq not installed. Run: pip install groq")
        return resume_path

    resume_text = _extract_resume_text(resume_path)

    prompt = TAILOR_PROMPT.format(
        job_title=job_title,
        company=company,
        job_description=job_description[:3000],
        resume_text=resume_text,
    )

    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.3,
        )
        response_text = response.choices[0].message.content.strip()

        json_match = re.search(r'\[.*\]', response_text, re.DOTALL)
        if not json_match:
            logger.warning("Groq returned no valid JSON for resume tailoring")
            return resume_path

        replacements = json.loads(json_match.group())

        if not replacements:
            logger.info(f"Resume already well-aligned with {job_title} at {company}")
            return resume_path

        safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
        safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')[:30]
        output_filename = f"Resume_{safe_title}_{safe_company}.docx"
        output_path = os.path.join(output_dir, output_filename)

        os.makedirs(output_dir, exist_ok=True)
        _apply_replacements(resume_path, replacements, output_path)
        logger.info(f"Tailored resume saved: {output_filename} ({len(replacements)} changes)")
        return output_path

    except Exception as e:
        logger.error(f"Resume tailoring failed: {e}")
        return resume_path
